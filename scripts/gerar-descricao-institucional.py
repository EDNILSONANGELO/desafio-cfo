# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Cores ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x16, 0x28)   # #0A1628 – fundo card
BLUE   = RGBColor(0x1D, 0x4E, 0xD8)   # #1D4ED8 – azul títulos
BLUE2  = RGBColor(0x4D, 0x8E, 0xF0)   # #4D8EF0 – azul secundário
GOLD   = RGBColor(0xC9, 0xA5, 0x20)   # #C9A520 – dourado
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x0F, 0x17, 0x2A)
GRAY   = RGBColor(0x47, 0x55, 0x69)
BODY   = RGBColor(0x33, 0x41, 0x55)
LIGHT  = RGBColor(0xE8, 0xED, 0xF5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_para_shading(para, fill_hex):
    """Aplica cor de fundo a um parágrafo."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val','single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color','000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run_color(run, color: RGBColor):
    run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(0)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    add_run_color(run, BLUE)
    run.font.name = 'Calibri'
    # linha separadora
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_before = Pt(0)
    border_p.paragraph_format.space_after  = Pt(6)
    set_para_shading(border_p, '1D4ED8')
    border_p.paragraph_format.space_before = Pt(0)
    border_p.paragraph_format.space_after  = Pt(8)
    border_run = border_p.add_run(' ')
    border_run.font.size = Pt(2)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    add_run_color(run, NAVY)
    run.font.name = 'Calibri'
    return p

def body(text, italic=False, color=None, space_after=Pt(5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = space_after
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.italic = italic
    add_run_color(run, color if color else BODY)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        add_run_color(r1, DARK)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
        add_run_color(r2, BODY)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        add_run_color(run, BODY)
    return p

def quote_box(text):
    """Caixa de citação estilizada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    set_para_shading(p, 'EFF6FF')
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = 'Calibri'
    add_run_color(run, BLUE)
    return p

def add_table_skills(rows_data):
    """Tabela de soft skills com cabeçalho azul."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # larguras
    for cell in table.columns[0].cells:
        cell.width = Cm(5.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(9.5)

    # cabeçalho
    hdr = table.rows[0].cells
    for i, txt in enumerate(['Soft Skill', 'Como é desenvolvida no jogo']):
        set_cell_shading(hdr[i], '1D4ED8')
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        add_run_color(run, WHITE)

    # linhas
    fills = ['FFFFFF', 'EFF6FF']
    for idx, (skill, desc) in enumerate(rows_data):
        row = table.add_row()
        fill = fills[idx % 2]
        # col 0 – skill
        set_cell_shading(row.cells[0], fill)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(skill)
        r0.bold = True
        r0.font.size = Pt(10.5)
        r0.font.name = 'Calibri'
        add_run_color(r0, NAVY)
        # col 1 – desc
        set_cell_shading(row.cells[1], fill)
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p1.add_run(desc)
        r1.font.size = Pt(10.5)
        r1.font.name = 'Calibri'
        add_run_color(r1, BODY)

    doc.add_paragraph()  # espaço após tabela

# ════════════════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════════════════
# Faixa de topo
p_top = doc.add_paragraph()
p_top.paragraph_format.space_before = Pt(0)
p_top.paragraph_format.space_after  = Pt(0)
set_para_shading(p_top, '0A1628')
r = p_top.add_run('  ')
r.font.size = Pt(4)

doc.add_paragraph()

# Título principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(6)
p_title.paragraph_format.space_after  = Pt(2)
r1 = p_title.add_run('Arena Contábil')
r1.bold = True
r1.font.size = Pt(28)
r1.font.name = 'Calibri'
add_run_color(r1, NAVY)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(2)
r2 = p_sub.add_run('Simulador Empresarial para o Ensino de Ciências Contábeis')
r2.font.size = Pt(14)
r2.font.name = 'Calibri'
r2.italic = True
add_run_color(r2, BLUE)

p_prog = doc.add_paragraph()
p_prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_prog.paragraph_format.space_before = Pt(4)
p_prog.paragraph_format.space_after  = Pt(4)
r3 = p_prog.add_run('Programa AI First — UniFECAF')
r3.font.size = Pt(11)
r3.font.name = 'Calibri'
add_run_color(r3, GRAY)

# Faixa dourada
p_gold = doc.add_paragraph()
p_gold.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_gold.paragraph_format.space_before = Pt(4)
p_gold.paragraph_format.space_after  = Pt(4)
set_para_shading(p_gold, 'C9A520')
rg = p_gold.add_run('  ')
rg.font.size = Pt(3)

p_autor = doc.add_paragraph()
p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_autor.paragraph_format.space_before = Pt(6)
p_autor.paragraph_format.space_after  = Pt(14)
ra = p_autor.add_run('Sistema idealizado e desenvolvido por Prof. Ednilson Angelo — UniFECAF')
ra.font.size = Pt(10)
ra.font.name = 'Calibri'
ra.italic = True
add_run_color(ra, GRAY)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 — APRESENTAÇÃO
# ════════════════════════════════════════════════════════════════════════════
heading1('1. Apresentação do Projeto')

body(
    'O Arena Contábil é um simulador empresarial acadêmico desenvolvido para transformar a forma como as Ciências Contábeis são ensinadas no ensino superior. Trata-se de uma plataforma digital completa, construída com tecnologia moderna e design inovador, que coloca o aluno no papel de gestor de uma empresa real — tomando decisões financeiras, contábeis e estratégicas em tempo real, competindo com seus colegas de turma e vivenciando diretamente as consequências de cada escolha nos resultados do negócio.'
)
body(
    'O sistema foi idealizado e desenvolvido pelo Prof. Ednilson Angelo como resposta a um dos maiores desafios do ensino contábil contemporâneo: a distância entre o conteúdo teórico ensinado em sala de aula e a realidade prática que o mercado exige do profissional de Ciências Contábeis.'
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 — PROBLEMA
# ════════════════════════════════════════════════════════════════════════════
heading1('2. O Problema que o Arena Contábil Resolve')

body(
    'O ensino tradicional de Contabilidade enfrenta um paradoxo histórico: forma profissionais tecnicamente preparados para interpretar demonstrações financeiras, mas raramente os expõe ao processo de gerar essas demonstrações a partir de decisões reais de gestão. O aluno aprende a ler um Balanço Patrimonial, mas não vivencia o impacto que uma decisão de produção, de crédito ou de política de preços provoca no resultado da empresa.'
)
body(
    'Esse distanciamento entre teoria e prática resulta em egressos que dominam a técnica, mas encontram dificuldade em aplicá-la com autonomia, velocidade e segurança no ambiente corporativo.'
)

quote_box(
    'Atualmente, o mercado valoriza cada vez mais profissionais que sabem aplicar conhecimento na prática, tomar decisões e gerar resultados — e não apenas possuir um diploma.'
)

body(
    'O mercado exige profissionais capazes de analisar cenários, interpretar dados e tomar decisões estratégicas com base em informações contábeis e financeiras. O Arena Contábil foi criado para preencher exatamente essa lacuna — permitindo que os alunos vivenciem desafios reais de gestão, tomem decisões estratégicas, desenvolvam competências técnicas e comportamentais e se preparem para o ambiente corporativo de forma prática e dinâmica, ainda dentro da sala de aula.'
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 — METODOLOGIA UNIFECAF
# ════════════════════════════════════════════════════════════════════════════
heading1('3. Alinhamento com a Metodologia da UniFECAF')

body(
    'O Arena Contábil nasce diretamente alinhado à essência da proposta pedagógica da UniFECAF: ensinar o que realmente importa para o mercado.'
)
body(
    'A UniFECAF tem como marca registrada o foco em competências práticas, o desenvolvimento de habilidades aplicadas, a preparação profissional baseada em experiências reais e a aproximação permanente entre ensino e realidade corporativa. O Arena Contábil materializa esses princípios de forma concreta e mensurável — cada rodada do simulador é, na prática, uma aula de gestão empresarial real, vivida pelos alunos, não apenas assistida.'
)
body(
    'Nesse sentido, o projeto não é apenas uma ferramenta tecnológica: é uma expressão pedagógica dos valores que a UniFECAF cultiva em todos os seus cursos — a crença de que aprender fazendo transforma carreiras.',
    color=BLUE
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 4 — COMO FUNCIONA
# ════════════════════════════════════════════════════════════════════════════
heading1('4. Como Funciona o Sistema')

body(
    'O Arena Contábil organiza a turma em grupos de alunos, cada grupo representando uma empresa concorrente dentro de um mercado simulado. No jogo, todas as empresas fabricam e vendem o mesmo produto — garrafas sustentáveis da marca EcoBottle — e competem entre si pelo melhor desempenho financeiro ao longo de diversas rodadas de decisão, cada uma representando um período contábil completo.'
)
body(
    'O professor atua como Game Master: cria as rodadas, define eventos econômicos que afetam o mercado — como crises econômicas, inflação, incentivos fiscais, escassez de insumos, alta do dólar e campanhas de sustentabilidade —, monitora os envios em tempo real e processa os resultados ao final de cada período. Os alunos não sabem qual evento será aplicado antes de enviarem suas decisões, simulando a imprevisibilidade real do ambiente de negócios.'
)

heading2('Decisões tomadas a cada rodada:')
for item in [
    ('Produção: ', 'quantidade a produzir, gestão de funcionários, contratações, demissões e expansão da capacidade produtiva via aquisição de máquinas;'),
    ('Materiais: ', 'compra de matérias-primas (plástico, tampa, embalagem e rótulo), com controle de estoque e custo de insumos;'),
    ('Vendas por Região: ', 'estratégia de distribuição em até cinco regiões geográficas, com definição de quantidade, preço de venda e investimento em marketing por mercado;'),
    ('Marketing: ', 'investimento em inserções de marketing e política de descontos comerciais para ampliar a demanda;'),
    ('Máquinas: ', 'compra de equipamentos para aumentar a capacidade produtiva;'),
    ('Finanças: ', 'tomada de empréstimos, definição do prazo médio de pagamento a fornecedores e do prazo médio de recebimento de clientes.'),
]:
    bullet(item[1], bold_prefix=item[0])

heading2('Resultados gerados automaticamente após cada rodada:')
for item in [
    'Demonstração do Resultado do Exercício (DRE) — regime de competência;',
    'Balanço Patrimonial — com Ativo, Passivo e Patrimônio Líquido, respeitando o princípio da continuidade contábil;',
    'Demonstração dos Fluxos de Caixa (DFC);',
    '13 indicadores financeiros: Liquidez Corrente, Seca e Imediata; Margem Bruta, Operacional e Líquida; ROA; ROE; PME; PMR; PMP; Ciclo Financeiro e Compras;',
    'Score ponderado por desempenho financeiro, gerando ranking competitivo em tempo real;',
    'Medalhas e reconhecimentos por categorias de excelência: melhor CFO, maior lucratividade, melhor liquidez, maior crescimento, campeão de vendas e rating AAA.',
]:
    bullet(item)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 5 — GAMIFICAÇÃO
# ════════════════════════════════════════════════════════════════════════════
heading1('5. Gamificação, Competição e Engajamento')

body(
    'Um dos diferenciais mais impactantes do Arena Contábil é a forma como a gamificação e a competição entre grupos transformam a dinâmica da sala de aula.'
)
body(
    'Em vez de assistir a uma aula expositiva, os alunos jogam. Em vez de resolver exercícios abstratos, eles gerenciam empresas. Em vez de esperar a próxima prova, eles acompanham seu ranking em tempo real e buscam estratégias para subir de posição na próxima rodada.'
)
body(
    'O sistema de ranking, medalhas, dashboards interativos e relatórios gerenciais cria um ciclo de engajamento natural: o aluno quer saber por que sua empresa foi mal, analisa os indicadores, ajusta a estratégia e testa as hipóteses na rodada seguinte — reproduzindo, em ambiente seguro e pedagógico, o ciclo de aprendizado por tentativa, análise e correção que define a prática profissional de alto desempenho.'
)
body(
    'A competição saudável entre grupos amplia ainda mais esse efeito: quando os colegas estão na frente no ranking, a motivação para melhorar é intrínseca, genuína e contínua — algo que nenhum sistema de avaliação tradicional consegue produzir com a mesma intensidade.'
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 6 — TECNOLOGIA
# ════════════════════════════════════════════════════════════════════════════
heading1('6. Tecnologia e Inovação Aplicada ao Ensino')

body(
    'O Arena Contábil representa uma aplicação direta da tecnologia e da inteligência artificial no ensino superior. O motor de simulação utiliza algoritmos que modelam o comportamento do mercado e reagem às decisões coletivas dos grupos — o preço médio de mercado, por exemplo, é calculado dinamicamente a partir dos preços praticados por todas as empresas simultaneamente. Isso significa que nenhuma rodada é igual à anterior: o comportamento do mercado emerge das decisões de todos, criando um ambiente competitivo vivo, imprevisível e pedagogicamente rico.'
)

heading2('Stack tecnológica utilizada:')
for item in [
    ('Frontend: ', 'Next.js 15, TypeScript e Tailwind CSS — interface responsiva, acessível em qualquer dispositivo, sem instalação;'),
    ('Backend: ', 'API Routes com banco de dados PostgreSQL hospedado na Supabase — escalável, seguro e rastreável;'),
    ('Autenticação: ', 'sistema JWT com controle de acesso por perfil (Professor, Aluno e Master);'),
    ('Motor de simulação: ', 'engine proprietária que calcula DRE, BP, DFC e 13 indicadores a partir de cada conjunto de decisões;'),
    ('Exportação: ', 'relatórios em PDF e planilhas Excel para uso acadêmico, avaliativo e de portfólio;'),
    ('Interface: ', 'design fintech com dashboard gerencial, gráficos de market share, evolução de lucro e ranking animado.'),
]:
    bullet(item[1], bold_prefix=item[0])

body(
    'Além disso, o sistema oferece ao aluno uma DRE prévia calculada em tempo real enquanto ele preenche suas decisões — permitindo que visualize o impacto financeiro de cada escolha antes de confirmar o envio, desenvolvendo o pensamento analítico e a tomada de decisão baseada em dados de forma imediata e intuitiva.'
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 7 — COMPETÊNCIAS
# ════════════════════════════════════════════════════════════════════════════
heading1('7. Competências Desenvolvidas')

heading2('7.1 Competências Técnicas — Hard Skills')

body('O Arena Contábil trabalha diretamente com o conteúdo das principais disciplinas do curso de Ciências Contábeis:')
for item in [
    'Elaboração e interpretação de DRE, Balanço Patrimonial e DFC;',
    'Cálculo e análise de 13 indicadores financeiros de liquidez, rentabilidade, atividade e estrutura patrimonial;',
    'Gestão do capital de giro: PME, PMR, PMP e Ciclo Financeiro;',
    'Formação do custo de produção e do preço de venda;',
    'Análise de margem bruta, operacional e líquida;',
    'Gestão da folha salarial, encargos trabalhistas e custos fixos;',
    'Decisões de financiamento: empréstimos, prazo de pagamento e recebimento;',
    'Análise de market share e posicionamento competitivo;',
    'Aplicação do regime de competência, princípio da continuidade e equação patrimonial;',
    'Leitura e interpretação de dashboards gerenciais e relatórios financeiros.',
]:
    bullet(item)

heading2('7.2 Competências Comportamentais — Soft Skills')

body('Além das competências técnicas, o ambiente competitivo e colaborativo do Arena Contábil desenvolve ativamente as soft skills mais valorizadas pelo mercado corporativo:')

skills = [
    ('Liderança', 'Cada grupo precisa definir e coordenar sua estratégia coletiva'),
    ('Trabalho em equipe', 'As decisões são tomadas em grupo — o resultado reflete a colaboração de todos'),
    ('Comunicação', 'Os grupos precisam argumentar, alinhar posições e defender suas estratégias'),
    ('Pensamento crítico', 'Cada rodada exige análise dos resultados anteriores para ajuste da próxima estratégia'),
    ('Resolução de problemas', 'Eventos econômicos inesperados exigem adaptação rápida das decisões'),
    ('Tomada de decisão', 'Toda rodada tem prazo — decidir sob pressão é parte do jogo'),
    ('Negociação', 'Definir preços, margens e alocação de recursos envolve trade-offs permanentes'),
    ('Visão estratégica', 'Planejar múltiplas rodadas exige pensar além do resultado imediato'),
    ('Gestão do tempo', 'O prazo de envio das decisões é real — os grupos aprendem a trabalhar com urgência'),
    ('Inteligência emocional', 'Lidar com quedas no ranking e reverter resultados negativos exige resiliência'),
    ('Análise de cenários', 'Os alunos projetam cenários e avaliam riscos antes de cada rodada'),
    ('Capacidade analítica', 'A interpretação de 13 indicadores e três demonstrações exige raciocínio quantitativo apurado'),
    ('Colaboração entre equipes', 'A turma aprende coletivamente com os resultados de todas as empresas'),
    ('Adaptação a mudanças', 'Eventos econômicos variáveis simulam a instabilidade real do ambiente de negócios'),
    ('Responsabilidade na gestão', 'As consequências das decisões são imediatas, visíveis e mensuráveis'),
]
add_table_skills(skills)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 8 — RESULTADOS PEDAGÓGICOS
# ════════════════════════════════════════════════════════════════════════════
heading1('8. Resultados Pedagógicos Práticos')

for (titulo, texto) in [
    ('Engajamento genuíno: ',
     'a dinâmica competitiva e o ranking em tempo real criam um ambiente de alta motivação — os alunos acompanham sua posição, analisam os resultados e buscam ativamente melhorar seu desempenho a cada rodada, comportamento raramente observado em aulas expositivas tradicionais.'),
    ('Compreensão integrada da contabilidade: ',
     'ao tomar decisões em seis áreas interdependentes e ver o resultado consolidado nas demonstrações contábeis, o aluno compreende pela primeira vez de forma visceral como as áreas de uma empresa se conectam.'),
    ('Retenção de longo prazo: ',
     'o aprendizado por experiência e erro gera memória duradoura. O aluno que perdeu posições no ranking porque praticou um preço acima do mercado, ou que teve caixa negativo por excesso de estoque, dificilmente esquecerá esses conceitos após a formatura.'),
    ('Desenvolvimento de análise crítica: ',
     'a cada rodada processada, o aluno recebe não apenas o resultado, mas os indicadores que explicam por que sua empresa foi bem ou mal — desenvolvendo a capacidade de leitura analítica que é o diferencial do profissional contábil de alto desempenho.'),
    ('Preparação real para o mercado: ',
     'ao término do jogo, o aluno terá elaborado múltiplas DREs, interpretado Balanços Patrimoniais, gerenciado capital de giro, analisado índices financeiros e tomado decisões estratégicas sob pressão competitiva — experiências que normalmente só viriam após anos de prática profissional.'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    r_bold = p.add_run(titulo)
    r_bold.bold = True
    r_bold.font.size = Pt(11)
    r_bold.font.name = 'Calibri'
    add_run_color(r_bold, NAVY)
    r_rest = p.add_run(texto)
    r_rest.font.size = Pt(11)
    r_rest.font.name = 'Calibri'
    add_run_color(r_rest, BODY)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 9 — PRESENCIAL E EAD
# ════════════════════════════════════════════════════════════════════════════
heading1('9. Aplicação nos Cursos Presenciais e EAD')

body(
    'O Arena Contábil foi projetado para funcionar plenamente em todos os formatos pedagógicos adotados pela educação superior contemporânea:'
)
for item in [
    ('Presencial: ', 'os grupos trabalham em sala de aula, com o professor mediando as rodadas em tempo real. A dinâmica de competição presencial cria um ambiente de energia e engajamento único;'),
    ('Semipresencial/Híbrido: ', 'as decisões são tomadas remotamente pelos grupos, com o professor processando os resultados em aula presencial — ideal para metodologias ativas híbridas;'),
    ('EAD: ', 'a plataforma é 100% online e responsiva. Professor e alunos interagem pelo sistema a qualquer hora e de qualquer dispositivo. O tracker de envios em tempo real garante ao professor visibilidade total mesmo à distância.'),
]:
    bullet(item[1], bold_prefix=item[0])

body(
    'O sistema suporta múltiplas turmas simultâneas, com dados completamente isolados entre elas. Pode ser aplicado com aproveitamento direto nas disciplinas de Contabilidade Gerencial, Análise das Demonstrações Contábeis, Custos e Formação de Preço, Finanças Empresariais, Contabilidade de Custos e Gestão Estratégica — funcionando tanto como ferramenta de avaliação quanto como metodologia ativa central da disciplina.'
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 10 — RELEVÂNCIA ESTRATÉGICA
# ════════════════════════════════════════════════════════════════════════════
heading1('10. Relevância Estratégica para a UniFECAF e o Programa AI First')

body(
    'O Arena Contábil materializa, de forma concreta e aplicada, os três pilares que definem a educação superior de excelência no século XXI:'
)
for item in [
    ('Inovação tecnológica no ensino: ', 'a plataforma utiliza tecnologias de ponta e uma engine de simulação proprietária, demonstrando que é possível construir ferramentas pedagógicas de alto impacto com a mesma qualidade técnica exigida pela indústria de tecnologia.'),
    ('Aprendizagem baseada em experiência: ', 'o sistema substitui a aula expositiva tradicional por um ambiente de aprendizado imersivo, orientado por dados, no qual o erro é pedagógico e o feedback é imediato — alinhado às melhores práticas internacionais de metodologias ativas.'),
    ('Preparação real para o mercado: ', 'em um cenário no qual os escritórios contábeis e as áreas financeiras das empresas já utilizam dashboards, indicadores em tempo real e análise de dados como rotina, o Arena Contábil prepara o egresso da UniFECAF para operar nesse ambiente com naturalidade, competência e confiança.'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    r_bold = p.add_run(item[0])
    r_bold.bold = True
    r_bold.font.size = Pt(11)
    r_bold.font.name = 'Calibri'
    add_run_color(r_bold, NAVY)
    r_rest = p.add_run(item[1])
    r_rest.font.size = Pt(11)
    r_rest.font.name = 'Calibri'
    add_run_color(r_rest, BODY)

body(
    'O projeto representa também um diferencial institucional relevante: uma ferramenta pedagógica proprietária, desenvolvida internamente por um professor da própria UniFECAF, que posiciona a instituição como referência em inovação acadêmica aplicada às Ciências Contábeis — capaz de valorizar a marca, fortalecer o programa de inovação e gerar reconhecimento junto a entidades como o CRC, o CFC e a comunidade acadêmica contábil brasileira.'
)

# ════════════════════════════════════════════════════════════════════════════
# CONSIDERAÇÕES FINAIS
# ════════════════════════════════════════════════════════════════════════════
heading1('Considerações Finais')

body(
    'O Arena Contábil não é apenas um jogo. É um laboratório contábil digital — um ambiente em que a teoria encontra a prática, em que a decisão tem consequência, em que o erro ensina e o acerto motiva.'
)
body(
    'É uma resposta concreta ao que o ensino superior contábil mais precisa: experiências que preparem o aluno não apenas para passar em uma prova, mas para enfrentar com segurança o primeiro dia de trabalho em um escritório, numa diretoria financeira ou na gestão de um negócio próprio.'
)

# Faixa final
p_final = doc.add_paragraph()
p_final.paragraph_format.space_before = Pt(16)
p_final.paragraph_format.space_after  = Pt(0)
set_para_shading(p_final, '0A1628')
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_final.add_run('Arena Contábil  |  Business Accounting Simulator  |  UniFECAF')
rf.font.size = Pt(10)
rf.font.name = 'Calibri'
rf.bold = True
add_run_color(rf, WHITE)

p_assinatura = doc.add_paragraph()
p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_assinatura.paragraph_format.space_before = Pt(6)
p_assinatura.paragraph_format.space_after  = Pt(0)
rs = p_assinatura.add_run('Sistema idealizado e desenvolvido por Prof. Ednilson Angelo — UniFECAF')
rs.font.size = Pt(9)
rs.font.name = 'Calibri'
rs.italic = True
add_run_color(rs, GRAY)

# ── Salvar ────────────────────────────────────────────────────────────────
output = 'C:/Users/ednil/desafio-cfo/Arena-Contabil-Descricao-Institucional.docx'
doc.save(output)
print(f'Documento salvo: {output}')
